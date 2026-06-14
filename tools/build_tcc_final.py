from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
import re
from pathlib import Path


OUT = Path("TCC_Portal_Digital_Comunidade_CC2026_FINAL.docx")


TITLE = "Portal Digital da Comunidade: Tecnologia Educacional para Memória, Saberes do Campo e Inclusão Digital"
AUTHORS = "Kauemberg Azevedo de Sousa; João Lucas da Silva Ferreira"

RESUMO = (
    "O projeto Portal Digital da Comunidade desenvolveu uma plataforma web educacional para registrar, organizar e divulgar "
    "memórias, saberes do campo, produção agrícola, biodiversidade, plantas medicinais e reportagens escolares vinculadas à "
    "EEMPC Francisco Araújo Barros, no Assentamento Lagoa do Mineiro, em Itarema, Ceará. A pesquisa foi realizada por meio "
    "de abordagem qualitativa, aplicada e participativa, articulando levantamento documental, pesquisa de campo, curadoria "
    "de conteúdos, desenvolvimento frontend com HTML5, CSS3 e JavaScript, publicação no GitHub Pages e preparação de um "
    "assistente educacional com inteligência artificial. Como resultado, foi criado um portal responsivo, acessível e "
    "modular, com páginas sobre história, memória, produção agrícola, plantas nativas, plantas medicinais, mudanças climáticas, "
    "educação do campo e IA na educação. A iniciativa fortalece a identidade camponesa, amplia a inclusão digital e transforma "
    "a escola em produtora de tecnologia social para preservar e compartilhar conhecimentos locais."
)

PALAVRAS_CHAVE = "Portal digital; Educação do campo; Memória comunitária; Inclusão digital; Inteligência artificial."

ABSTRACT = (
    "The Digital Community Portal project developed an educational web platform to record, organize and share memories, rural "
    "knowledge, agricultural production, biodiversity, medicinal plants and school reports connected to EEMPC Francisco Araújo "
    "Barros, located in the Lagoa do Mineiro settlement, Itarema, Ceará, Brazil. The research followed a qualitative, applied "
    "and participatory approach, combining documentary review, field research, content curation, frontend development with "
    "HTML5, CSS3 and JavaScript, GitHub Pages publication and preparation for an educational assistant powered by artificial "
    "intelligence. As a result, a responsive, accessible and modular portal was created, including pages about history, memory, "
    "agricultural production, native plants, medicinal plants, climate change, rural education and AI in education. The project "
    "strengthens peasant identity, expands digital inclusion and positions the school as a producer of social technology for "
    "preserving and sharing local knowledge."
)

KEYWORDS = "Digital portal; Rural education; Community memory; Digital inclusion; Artificial intelligence."

INTRODUCAO = (
    "O Brasil possui grande diversidade cultural, histórica e territorial, mas parte significativa desse patrimônio permanece "
    "pouco registrada em ambientes digitais, especialmente nas comunidades rurais. No Assentamento Lagoa do Mineiro, em Itarema, "
    "Ceará, a história de luta pela terra, a memória dos moradores, os saberes agrícolas, o conhecimento sobre plantas nativas "
    "e medicinais e a experiência cotidiana da escola do campo formam um acervo vivo, transmitido principalmente pela oralidade. "
    "Quando esse acervo não é documentado, corre o risco de desaparecer com o tempo, enfraquecendo a identidade coletiva e "
    "dificultando que as novas gerações reconheçam o valor científico, cultural e social do próprio território. "
    "A EEMPC Francisco Araújo Barros, por ser uma escola do campo inserida nesse território, assume papel estratégico na "
    "produção de conhecimento contextualizado. A educação do campo defende que a aprendizagem deve dialogar com a vida real "
    "dos sujeitos, com a cultura local, com o trabalho, com a memória e com o desenvolvimento sustentável das comunidades. "
    "Nesse sentido, a tecnologia não deve ser entendida apenas como ferramenta técnica, mas como meio de cidadania, expressão "
    "cultural e democratização do acesso à informação. "
    "O problema investigado parte da constatação de que a comunidade possui muitos conhecimentos importantes, mas pouca "
    "presença organizada na internet. Informações sobre a origem do assentamento, a atuação dos moradores, as práticas agrícolas, "
    "as espécies vegetais da região e os projetos escolares estavam dispersas, dependentes de relatos orais ou registradas em "
    "arquivos locais pouco acessíveis. Ao mesmo tempo, estudantes do curso técnico em informática dominam ferramentas digitais "
    "capazes de transformar esse cenário. "
    "A justificativa do projeto está na necessidade de unir tecnologia, educação, memória e impacto social. Ao criar um portal "
    "digital, os estudantes exercitam programação, design, acessibilidade, escrita jornalística, curadoria de dados, pesquisa "
    "participativa e responsabilidade ética no uso de informações. A comunidade, por sua vez, passa a ter um espaço público para "
    "consultar, compartilhar e valorizar sua própria história. "
    "A fundamentação teórica articula quatro eixos. O primeiro é a memória coletiva, compreendida como construção social que "
    "mantém vivos os vínculos de pertencimento. O segundo é o patrimônio cultural imaterial, formado por práticas, saberes, "
    "expressões e modos de fazer reconhecidos pela comunidade. O terceiro é a educação libertadora e contextualizada, inspirada "
    "em Paulo Freire, que parte da realidade concreta dos educandos. O quarto é a inclusão digital, entendida como capacidade de "
    "usar tecnologias para participação social, produção de conhecimento e redução de desigualdades. "
    "Dessa forma, o Portal Digital da Comunidade responde a uma demanda educacional e social: criar uma tecnologia simples, "
    "acessível e sustentável, desenvolvida pelos próprios estudantes, para fortalecer a memória, a pesquisa escolar e a presença "
    "digital de uma comunidade rural cearense. "
    "A proposta também se justifica pela possibilidade de continuidade. Diferentemente de uma atividade pontual, o portal pode "
    "ser alimentado por novas turmas, receber fotografias, entrevistas, reportagens, mapas, calendários e materiais pedagógicos. "
    "Assim, o produto final não se limita a demonstrar domínio técnico em programação, mas inaugura uma infraestrutura escolar "
    "permanente para pesquisa, comunicação comunitária e aprendizagem interdisciplinar."
)

OBJ_GERAL = (
    "Desenvolver e documentar um portal web educacional, acessível e responsivo, capaz de registrar a memória comunitária, "
    "valorizar saberes do campo, divulgar pesquisas escolares e preparar recursos interativos com inteligência artificial para "
    "estudantes, professores e moradores."
)

OBJ_ESPECIFICOS = (
    "Mapear conteúdos históricos, culturais, agrícolas e ambientais relevantes para a comunidade; organizar páginas digitais "
    "sobre história, memória, produção agrícola, plantas nativas, plantas medicinais e reportagens; aplicar HTML, CSS e "
    "JavaScript na construção de uma plataforma responsiva; implementar espaços para imagens, notícias, eventos e conteúdo "
    "dinâmico; preparar arquitetura segura para integração futura com assistente educacional de IA; validar o portal como "
    "instrumento de aprendizagem, preservação cultural e inclusão digital."
)

METODOS = (
    "A pesquisa caracteriza-se como qualitativa, aplicada e participativa, com elementos de pesquisa-ação e desenvolvimento "
    "tecnológico. É qualitativa porque analisa memórias, narrativas, saberes e práticas sociais que não se reduzem a números; "
    "aplicada porque resulta em um produto concreto, o Portal Digital da Comunidade; e participativa porque envolve estudantes, "
    "professores e comunidade na seleção, organização e validação dos conteúdos. "
    "O universo do projeto é a EEMPC Francisco Araújo Barros e as comunidades relacionadas ao Assentamento Lagoa do Mineiro, em "
    "Itarema, Ceará. Os participantes atuaram como pesquisadores escolares, desenvolvedores, curadores de conteúdo e validadores "
    "das informações. Foram considerados registros históricos, relatos comunitários, dados agrícolas, informações sobre espécies "
    "nativas e medicinais, reportagens escolares e materiais técnicos produzidos durante o desenvolvimento do portal. "
    "A construção tecnológica utilizou HTML5 para estruturar as páginas, CSS3 para identidade visual e responsividade, JavaScript "
    "moderno para menu, carrossel, formulário, carregamento de JSON e assistente educacional, além de GitHub Pages para publicação "
    "do frontend. O projeto também estruturou um backend Node.js opcional para conexão segura com API de IA, mantendo chaves fora "
    "do navegador. "
    "As etapas executadas foram: diagnóstico da ausência de registros digitais organizados; definição da arquitetura do portal; "
    "criação das páginas principais; organização de conteúdo dinâmico em JSON; preparação de placeholders para fotos reais dos "
    "alunos; implementação de acessibilidade; publicação do site; revisão visual; correção de links, acentuação e padronização "
    "do layout; e documentação técnica. A análise dos resultados considerou coerência pedagógica, funcionamento técnico, "
    "acessibilidade, potencial de uso comunitário e alinhamento com o Ceará Científico."
)

RESULTADOS = (
    "O principal resultado foi a construção de um portal educacional publicado no GitHub Pages, disponível em endereço público, "
    "com identidade visual própria e navegação organizada por temas. A página inicial apresenta o projeto, rolo de fotos, atalhos "
    "para áreas de pesquisa, notícias, eventos e chamada para o assistente de IA. As páginas internas abordam história, memória, "
    "produção agrícola, plantação nativa, plantas medicinais, produção de mudas, educação do campo, mudanças climáticas e IA na "
    "educação. "
    "No campo técnico, o portal avançou de um conjunto de páginas estáticas para uma arquitetura modular. O JavaScript foi separado "
    "em componentes de menu, formulário, carrossel, imagens, conteúdo dinâmico e serviços de IA. O conteúdo leve passou a ser "
    "organizado em arquivo JSON, facilitando atualização de notícias, eventos e listas sem alterar toda a estrutura do site. O "
    "CSS principal foi padronizado para manter unidade visual entre as páginas, com responsividade, cards, seções, rodapé comum e "
    "melhor experiência em dispositivos móveis. "
    "No campo pedagógico, o portal transformou a comunidade em objeto legítimo de investigação. A história do assentamento, os "
    "saberes sobre plantas, a produção agrícola e as práticas culturais passaram a ser tratados como conhecimento escolar e "
    "científico. Isso fortalece a autoestima dos estudantes e aproxima o curso técnico em informática das necessidades reais do "
    "território. "
    "Outro resultado relevante foi a preparação para inteligência artificial educacional. O projeto não expõe chaves no frontend; "
    "ao contrário, prevê backend seguro para responder dúvidas, recomendar páginas, gerar quizzes e apoiar pesquisas. Essa escolha "
    "mostra maturidade técnica e preocupação ética. "
    "A discussão evidencia que a inovação do projeto não está apenas no uso de tecnologia, mas na combinação entre tecnologia, "
    "memória, cultura regional, educação do campo e inclusão digital. O portal atua como produto técnico, ferramenta pedagógica e "
    "tecnologia social, podendo ser ampliado continuamente por novas turmas."
)

CONSIDERACOES = (
    "O Portal Digital da Comunidade alcançou seu objetivo ao criar uma plataforma web educacional capaz de reunir memória, "
    "cultura, produção agrícola, biodiversidade e inovação tecnológica em um único ambiente acessível. O projeto demonstra que "
    "estudantes da escola do campo podem produzir tecnologia relevante, conectada à realidade local e com potencial de impacto "
    "social. "
    "A experiência permitiu desenvolver competências de programação, design, escrita, pesquisa, curadoria, acessibilidade e "
    "trabalho colaborativo. Também mostrou que a internet pode ser usada para combater a invisibilidade de comunidades rurais, "
    "registrando saberes que muitas vezes permanecem restritos à oralidade. "
    "Como continuidade, recomenda-se ampliar a coleta de depoimentos, inserir fotos reais dos alunos, publicar novas reportagens, "
    "validar o portal com moradores e ativar o backend de IA. Assim, o projeto pode evoluir de protótipo escolar para acervo "
    "digital permanente da comunidade, fortalecendo o Ceará Científico como espaço de ciência, cidadania e transformação social. "
    "Seu maior valor está em provar que inovação também nasce do território."
)

REFERENCIAS = [
    "CEARÁ. Secretaria da Educação do Estado do Ceará. Regulamento Ceará Científico: Mais Solidário, Mais Cooperativo 2026. Fortaleza: SEDUC, 2026.",
    "BRASIL. Decreto nº 7.352, de 4 de novembro de 2010. Dispõe sobre a política de educação do campo e o Programa Nacional de Educação na Reforma Agrária.",
    "BRASIL. Decreto nº 3.551, de 4 de agosto de 2000. Institui o registro de bens culturais de natureza imaterial.",
    "BRASIL. Lei nº 13.709, de 14 de agosto de 2018. Lei Geral de Proteção de Dados Pessoais.",
    "CETIC.BR. TIC Educação: pesquisa sobre o uso das tecnologias de informação e comunicação nas escolas brasileiras.",
    "EMBRAPA. Tecnologias sociais, agricultura familiar e convivência com o semiárido. Brasília: Embrapa.",
    "FREIRE, Paulo. Pedagogia do Oprimido. Rio de Janeiro: Paz e Terra, 1987.",
    "HALBWACHS, Maurice. A memória coletiva. São Paulo: Centauro, 2006.",
    "IBGE. Censo Agropecuário e informações territoriais do Brasil. Rio de Janeiro: IBGE.",
    "UNESCO. Convenção para a Salvaguarda do Patrimônio Cultural Imaterial. Paris, 2003.",
    "W3C. Web Content Accessibility Guidelines (WCAG). World Wide Web Consortium.",
    "GITHUB. GitHub Pages Documentation. Disponível em: https://docs.github.com/pages.",
    "PORTAL DA COMUNIDADE. Site público do projeto. Disponível em: https://paulocunha1009.github.io/Portal-da-Comunidade/index.html."
]


def word_count(text):
    return len(re.findall(r"\b[\wÀ-ÿ]+\b", text))


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(9)
    return p


def add_table(doc, rows, widths=None, header=True):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Cm(width)
    for r_idx, row_data in enumerate(rows):
        for c_idx, value in enumerate(row_data):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_text(cell, str(value), bold=(header and r_idx == 0), align=WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT)
            if header and r_idx == 0:
                set_cell_shading(cell, "EAF2EA")
    return table


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    p.add_run(text)
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    p.style = "Normal"
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run(text)
    return p


def add_labelled(doc, label, text):
    p = doc.add_paragraph()
    p.style = "Normal"
    r = p.add_run(label + " ")
    r.bold = True
    p.add_run(text)
    return p


def apply_styles(doc):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(8)
    for name, size, color in [
        ("Heading 1", 15, "174726"),
        ("Heading 2", 13, "246B3A"),
        ("Heading 3", 12, "246B3A"),
    ]:
        st = styles[name]
        st.font.name = "Arial"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(12)
        st.paragraph_format.space_after = Pt(6)


def build():
    doc = Document()
    apply_styles(doc)

    # Capa
    for line in [
        "SECRETARIA DA EDUCAÇÃO DO ESTADO DO CEARÁ – SEDUC/CE",
        "CEARÁ CIENTÍFICO 2026",
        "EEMPC FRANCISCO ARAÚJO BARROS",
        "CURSO TÉCNICO EM INFORMÁTICA",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        r.bold = True
        r.font.name = "Arial"
        r.font.size = Pt(11)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(TITLE.upper())
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(16)
    r.font.color.rgb = RGBColor(23, 71, 38)

    for line in [
        "Autores: " + AUTHORS,
        "Professor orientador: Paulo Henrique Gomes Cordeiro da Cunha",
        "Município/CREDE: Itarema – CREDE 3",
        "Categoria: II – Ensino Médio: Ações Afirmativas e CEJAs EM | Escola do Campo",
        "Área de pesquisa: Robótica, Automação e Aplicação das TIC",
        "Eixos sugeridos: Educação para a Cidadania; Comunicação, Mídia e Democracia",
        "Produto: portal web educacional publicado no GitHub Pages",
        "URL: https://paulocunha1009.github.io/Portal-da-Comunidade/index.html",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        r.font.name = "Arial"
        r.font.size = Pt(11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Itarema – Ceará\n2026").bold = True

    doc.add_page_break()

    add_heading(doc, "Enquadramento no Ceará Científico 2026", 1)
    add_body(doc, (
        "Conforme o Regulamento Ceará Científico 2026, a inscrição deve indicar categoria, área de pesquisa e, quando pertinente, "
        "eixos temáticos. O enquadramento mais adequado para o projeto é a área Robótica, Automação e Aplicação das TIC, prevista "
        "no Anexo I para projetos relacionados ao desenvolvimento de softwares e à criação, compreensão, utilização e aplicação "
        "crítica, significativa, reflexiva e ética das Tecnologias da Informação e da Comunicação. O produto central desta pesquisa "
        "é justamente uma tecnologia digital desenvolvida por estudantes: um portal web com HTML, CSS, JavaScript, conteúdo dinâmico, "
        "acessibilidade, publicação em GitHub Pages e arquitetura para integração segura com inteligência artificial."
    ))
    add_body(doc, (
        "A categoria indicada é II – Ensino Médio: Ações Afirmativas e CEJAs EM, pois o edital inclui nessa categoria as Escolas "
        "do Campo. Como eixos temáticos sugeridos, o projeto dialoga especialmente com Educação para a Cidadania e Comunicação, "
        "Mídia e Democracia, já que utiliza tecnologia para fortalecer participação estudantil, memória comunitária, acesso público "
        "à informação, convivência democrática e valorização do território rural."
    ))

    add_heading(doc, "Conferência dos elementos do projeto", 1)
    section_counts = {
        "Título": word_count(TITLE),
        "Autores": word_count(AUTHORS),
        "Resumo": word_count(RESUMO),
        "Abstract/Resumen": word_count(ABSTRACT),
        "Introdução": word_count(INTRODUCAO),
        "Objetivo geral": word_count(OBJ_GERAL),
        "Objetivos específicos": word_count(OBJ_ESPECIFICOS),
        "Materiais e métodos": word_count(METODOS),
        "Resultados e discussão": word_count(RESULTADOS),
        "Considerações finais": word_count(CONSIDERACOES),
        "Referências bibliográficas": word_count(" ".join(REFERENCIAS)),
    }
    total_words = sum(section_counts.values())
    counts = [
        ["Tópico", "Mínimo", "Máximo", "Quantidade nesta versão"],
        ["Título", "1", "40", section_counts["Título"]],
        ["Autores", "1", "50", section_counts["Autores"]],
        ["Resumo", "100", "200", section_counts["Resumo"]],
        ["Palavras-chave", "1", "5", 5],
        ["Abstract/Resumen", "100", "200", section_counts["Abstract/Resumen"]],
        ["Keywords/Palavras clave", "1", "5", 5],
        ["Introdução", "480", "1000", section_counts["Introdução"]],
        ["Objetivo geral", "15", "100", section_counts["Objetivo geral"]],
        ["Objetivos específicos", "30", "200", section_counts["Objetivos específicos"]],
        ["Materiais e métodos", "200", "600", section_counts["Materiais e métodos"]],
        ["Resultados e discussão", "200", "800", section_counts["Resultados e discussão"]],
        ["Considerações finais", "150", "300", section_counts["Considerações finais"]],
        ["Referências bibliográficas", "10", "600", section_counts["Referências bibliográficas"]],
        ["Total de palavras", "-", "-", total_words],
        ["Imagens", "0 itens", "8 itens", "0 no DOCX; imagens disponíveis no portal"],
        ["Mídia", "link de vídeo", "1 a 5 min", "roteiro no Apêndice B"],
    ]
    add_table(doc, counts, widths=[6.8, 2.2, 2.2, 4.1])

    doc.add_page_break()

    add_heading(doc, "Título", 1)
    add_body(doc, TITLE)

    add_heading(doc, "Autores", 1)
    add_body(doc, AUTHORS)
    add_labelled(doc, "Orientador:", "Paulo Henrique Gomes Cordeiro da Cunha.")
    add_labelled(doc, "Instituição:", "EEMPC Francisco Araújo Barros, Itarema – CE.")

    add_heading(doc, "Resumo", 1)
    add_body(doc, RESUMO)

    add_heading(doc, "Palavras-chave", 1)
    add_body(doc, PALAVRAS_CHAVE)

    add_heading(doc, "Abstract", 1)
    add_body(doc, ABSTRACT)

    add_heading(doc, "Keywords", 1)
    add_body(doc, KEYWORDS)

    add_heading(doc, "1. Introdução", 1)
    for part in INTRODUCAO.split(". "):
        if part.strip():
            add_body(doc, part.strip() + ("" if part.endswith(".") else "."))

    add_heading(doc, "2. Objetivo geral", 1)
    add_body(doc, OBJ_GERAL)

    add_heading(doc, "3. Objetivos específicos", 1)
    for item in [
        "Mapear conteúdos históricos, culturais, agrícolas e ambientais relevantes para a comunidade.",
        "Organizar páginas digitais sobre história, memória, produção agrícola, plantas nativas, plantas medicinais e reportagens.",
        "Aplicar HTML, CSS e JavaScript na construção de uma plataforma responsiva e acessível.",
        "Implementar espaços para imagens, notícias, eventos e conteúdo dinâmico em JSON.",
        "Preparar arquitetura segura para integração futura com assistente educacional de IA.",
        "Validar o portal como instrumento de aprendizagem, preservação cultural e inclusão digital.",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)

    add_heading(doc, "4. Materiais e métodos", 1)
    for part in METODOS.split(". "):
        if part.strip():
            add_body(doc, part.strip() + ("" if part.endswith(".") else "."))

    add_heading(doc, "5. Resultados e discussão", 1)
    for part in RESULTADOS.split(". "):
        if part.strip():
            add_body(doc, part.strip() + ("" if part.endswith(".") else "."))

    add_heading(doc, "Síntese do produto desenvolvido", 2)
    add_table(doc, [
        ["Elemento do portal", "Resultado técnico-pedagógico"],
        ["Página inicial", "Apresenta o projeto, rolo de fotos, atalhos, notícias, eventos e chamada para IA."],
        ["História e memória", "Registra luta pela terra, identidade comunitária, depoimentos e saberes locais."],
        ["Produção agrícola", "Organiza dados de cultivo, práticas sustentáveis e relação com a economia familiar."],
        ["Plantas nativas e medicinais", "Valoriza biodiversidade, saberes tradicionais e pesquisa escolar responsável."],
        ["Reportagens", "Aprofunda educação do campo, mudanças climáticas e inteligência artificial na educação."],
        ["Arquitetura de IA", "Prevê backend seguro para respostas, quizzes e recomendações de conteúdo."],
    ], widths=[4.6, 10.7])

    add_heading(doc, "6. Considerações finais", 1)
    for part in CONSIDERACOES.split(". "):
        if part.strip():
            add_body(doc, part.strip() + ("" if part.endswith(".") else "."))

    add_heading(doc, "7. Referências bibliográficas", 1)
    for ref in REFERENCIAS:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(-0.5)
        p.paragraph_format.left_indent = Cm(0.5)
        p.add_run(ref)

    add_heading(doc, "Apêndice A — Link e evidências do produto", 1)
    add_labelled(doc, "Portal publicado:", "https://paulocunha1009.github.io/Portal-da-Comunidade/index.html")
    add_labelled(doc, "Repositório:", "github.com/paulocunha1009/Portal-da-Comunidade")
    add_labelled(doc, "Tecnologias:", "HTML5, CSS3, JavaScript, JSON, GitHub Pages, Node.js e backend opcional para IA.")
    add_labelled(doc, "Documentação técnica:", "DOCUMENTACAO_UNIFICADA.md e documentação do repositório local.")

    add_heading(doc, "Apêndice B — Roteiro para vídeo de apresentação", 1)
    add_body(doc, (
        "O vídeo de inscrição deve ser gravado na horizontal, em local silencioso e iluminado, conforme orientação do edital. "
        "Sugere-se duração entre 2 e 3 minutos, com participação dos dois estudantes."
    ))
    for item in [
        "Apresentar escola, autores, orientador, categoria, área de pesquisa e título.",
        "Explicar o problema: ausência de registros digitais organizados da memória e dos saberes da comunidade.",
        "Mostrar o portal publicado, destacando página inicial, história, memória, agricultura, plantas e reportagens.",
        "Explicar a metodologia: pesquisa participativa, curadoria de conteúdo, desenvolvimento web e validação.",
        "Apresentar resultados: site responsivo, documentação, conteúdo dinâmico, acessibilidade e arquitetura para IA.",
        "Concluir relacionando o projeto ao tema Ciência, Cidadania e Convivência Democrática: o conhecimento a serviço da vida coletiva.",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)

    doc.save(OUT)
    print(OUT.resolve())


if __name__ == "__main__":
    build()
